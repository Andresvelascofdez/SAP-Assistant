"""
Parseadores de archivos
Wiki Inteligente SAP IS-U
"""
import os
from typing import Optional, Dict, Any
from pathlib import Path
import logging

# Imports opcionales para parsers
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

logger = logging.getLogger(__name__)


class FileParser:
    """Parser para diferentes tipos de archivos"""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'parse_pdf',
        '.docx': 'parse_docx', 
        '.doc': 'parse_docx',
        '.md': 'parse_markdown',
        '.markdown': 'parse_markdown',
        '.html': 'parse_html',
        '.htm': 'parse_html',
        '.txt': 'parse_text'
    }
    
    def parse_file(self, file_path: str, content_type: str = None) -> Dict[str, Any]:
        """Parsear archivo y extraer texto"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Detectar tipo por extensión
        extension = Path(file_path).suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")
        
        parser_method = getattr(self, self.SUPPORTED_EXTENSIONS[extension])
        
        try:
            result = parser_method(file_path)
            result['file_type'] = extension
            result['file_size'] = os.path.getsize(file_path)
            return result
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {e}")
            raise ValueError(f"Failed to parse file: {str(e)}")
    
    def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parsear archivo PDF"""
        if not PDF_AVAILABLE:
            raise ValueError("PDF parsing not available. Install pdfminer.six")
        
        try:
            text = pdf_extract_text(file_path)
            return {
                'content': text.strip(),
                'metadata': {'pages': text.count('\f') + 1}
            }
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parsear archivo DOCX/DOC"""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX parsing not available. Install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extraer texto de párrafos
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extraer texto de tablas
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            
            content = '\n\n'.join(paragraphs)
            if tables_text:
                content += '\n\n--- TABLES ---\n\n' + '\n'.join(tables_text)
            
            return {
                'content': content,
                'metadata': {
                    'paragraphs': len(paragraphs),
                    'tables': len(doc.tables)
                }
            }
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    def parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """Parsear archivo Markdown"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Si markdown está disponible, convertir a HTML y extraer texto plano
            if HTML_AVAILABLE and 'markdown' in globals():
                html = markdown.markdown(content)
                soup = BeautifulSoup(html, 'html.parser')
                plain_text = soup.get_text()
                return {
                    'content': plain_text.strip(),
                    'metadata': {'original_format': 'markdown'}
                }
            else:
                # Fallback: devolver markdown raw
                return {
                    'content': content.strip(),
                    'metadata': {'original_format': 'markdown_raw'}
                }
        except Exception as e:
            raise ValueError(f"Failed to parse Markdown: {str(e)}")
    
    def parse_html(self, file_path: str) -> Dict[str, Any]:
        """Parsear archivo HTML"""
        if not HTML_AVAILABLE:
            raise ValueError("HTML parsing not available. Install beautifulsoup4")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remover scripts y estilos
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extraer texto
            text = soup.get_text()
            
            # Limpiar texto
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return {
                'content': text,
                'metadata': {'original_format': 'html'}
            }
        except Exception as e:
            raise ValueError(f"Failed to parse HTML: {str(e)}")
    
    def parse_text(self, file_path: str) -> Dict[str, Any]:
        """Parsear archivo de texto plano"""
        try:
            # Intentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Could not decode file with any supported encoding")
            
            return {
                'content': content.strip(),
                'metadata': {'encoding': encoding}
            }
        except Exception as e:
            raise ValueError(f"Failed to parse text file: {str(e)}")
    
    @classmethod
    def get_supported_extensions(cls) -> list:
        """Obtener lista de extensiones soportadas"""
        return list(cls.SUPPORTED_EXTENSIONS.keys())
    
    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """Verificar si el archivo es soportado"""
        extension = Path(file_path).suffix.lower()
        return extension in cls.SUPPORTED_EXTENSIONS
